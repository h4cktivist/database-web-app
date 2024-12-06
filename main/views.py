import os
import threading
import openpyxl
from django.http import JsonResponse
from docx import Document
from docx.enum.section import WD_ORIENT

from django.shortcuts import render, get_object_or_404, redirect, reverse
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.utils.timezone import now
from django.conf import settings
from django.db.models import Count, Value as V, CharField
from django.db.models.functions import Lower, Concat, Cast
from django.contrib import messages

from .models import Customers, Halls, Movies, Positions, SessionTypes, Staff, Sessions, Tickets, Sales
from .forms import CustomerForm, HallsForm, MoviesForm, PositionsForm, SessionTypesForm, StaffForm, SessionsForm, TicketsForm, SalesForm


def index(request):
    context = {
        'db_name': settings.DATABASES['default']['NAME'],
        'db_user': settings.DATABASES['default']['USER'],
        'db_host': settings.DATABASES['default']['HOST'],
        'db_port': settings.DATABASES['default']['PORT'],
    }
    return render(request, 'index.html', context)


def customers(request):
    sort_by = request.GET.get('sort_by', 'customer_id')
    sort_order = request.GET.get('sort_order', 'asc')

    columns = [
        {'field': 'customer_id', 'label': 'ID'},
        {'field': 'first_name', 'label': 'Имя'},
        {'field': 'last_name', 'label': 'Фамилия'},
        {'field': 'phone', 'label': 'Телефон'},
        {'field': 'email', 'label': 'Email'},
    ]
    for col in columns:
        if col['field'] == sort_by:
            col['sort_order'] = 'desc' if sort_order == 'asc' else 'asc'
        else:
            col['sort_order'] = 'asc'

    customer_list = Customers.objects.all().order_by(f"{'-' if sort_order == 'desc' else ''}{sort_by}")
    paginator = Paginator(customer_list, 25)
    page = request.GET.get('page', 1)
    try:
        customers = paginator.page(page)
    except PageNotAnInteger:
        customers = paginator.page(1)
    except EmptyPage:
        customers = paginator.page(paginator.num_pages)

    context = {
        'customers': customers,
        'columns': columns,
        'sort_by': sort_by,
        'sort_order': sort_order,
        'page_number': page,
    }
    return render(request, 'customers/customers.html', context)


def customer_detail(request, id):
    customer = get_object_or_404(Customers, pk=id)
    return render(request, 'customers/customer_detail.html', {'customer': customer})


def add_customer(request):
    if request.method == 'POST':
        form = CustomerForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('customers')
    else:
        form = CustomerForm()
    return render(request, 'add.html', {'form': form})


def edit_customer(request, id):
    customer = get_object_or_404(Customers, pk=id)
    if request.method == 'POST':
        form = CustomerForm(request.POST, instance=customer)
        if form.is_valid():
            form.save()
            return redirect('customers')
    else:
        form = CustomerForm(instance=customer)
    return render(request, 'edit.html', {'form': form, 'customer': customer})


def delete_customer(request, id):
    customer = get_object_or_404(Customers, pk=id)
    if request.method == 'POST':
        customer.delete()
        return redirect('customers')
    return render(request, 'delete.html', {'redirect_url': reverse('customers')})


def halls(request):
    sort_by = request.GET.get('sort_by', 'hall_id')
    sort_order = request.GET.get('sort_order', 'asc')

    columns = [
        {'field': 'hall_id', 'label': 'ID'},
        {'field': 'name', 'label': 'Название'},
        {'field': 'capacity', 'label': 'Вместимость'},
    ]
    for col in columns:
        if col['field'] == sort_by:
            col['sort_order'] = 'desc' if sort_order == 'asc' else 'asc'
        else:
            col['sort_order'] = 'asc'

    hall_list = Halls.objects.all().order_by(f"{'-' if sort_order == 'desc' else ''}{sort_by}")

    context = {
        'halls': hall_list,
        'columns': columns,
        'sort_by': sort_by,
        'sort_order': sort_order,
    }
    return render(request, 'halls/halls.html', context)


def hall_detail(request, id):
    hall = get_object_or_404(Halls, pk=id)
    return render(request, 'halls/hall_detail.html', {'hall': hall})


def add_hall(request):
    if request.method == 'POST':
        form = HallsForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('halls')
    else:
        form = HallsForm()
    return render(request, 'add.html', {'form': form})


def edit_hall(request, id):
    hall = get_object_or_404(Halls, pk=id)
    if request.method == 'POST':
        form = HallsForm(request.POST, instance=hall)
        if form.is_valid():
            form.save()
            return redirect('halls')
    else:
        form = HallsForm(instance=hall)
    return render(request, 'edit.html', {'form': form, 'hall': hall})


def delete_hall(request, id):
    hall = get_object_or_404(Halls, pk=id)
    if request.method == 'POST':
        hall.delete()
        return redirect('halls')
    return render(request, 'delete.html', {'redirect_url': reverse('halls')})


def movies(request):
    sort_by = request.GET.get('sort_by', 'movie_id')
    sort_order = request.GET.get('sort_order', 'asc')
    page_number = request.GET.get('page', 1)

    columns = [
        {'field': 'movie_id', 'label': 'ID'},
        {'field': 'title', 'label': 'Название'},
        {'field': 'genre', 'label': 'Жанр'},
        {'field': 'duration', 'label': 'Длительность'},
        {'field': 'age_restriction', 'label': 'Ограничение'},
        {'field': 'rating', 'label': 'Рейтинг'},
    ]
    for col in columns:
        if col['field'] == sort_by:
            col['sort_order'] = 'desc' if sort_order == 'asc' else 'asc'
        else:
            col['sort_order'] = 'asc'

    movie_list = Movies.objects.all().order_by(f"{'-' if sort_order == 'desc' else ''}{sort_by}")
    paginator = Paginator(movie_list, 25)
    try:
        movies = paginator.page(page_number)
    except PageNotAnInteger:
        movies = paginator.page(1)
    except EmptyPage:
        movies = paginator.page(paginator.num_pages)

    context = {
        'movies': movies,
        'columns': columns,
        'sort_by': sort_by,
        'sort_order': sort_order,
        'page_number': page_number,
    }
    return render(request, 'movies/movies.html', context)


def movie_detail(request, movie_id):
    movie = get_object_or_404(Movies, movie_id=movie_id)
    return render(request, 'movies/movie_detail.html', {'movie': movie})


def add_movie(request):
    if request.method == 'POST':
        form = MoviesForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('movies')
    else:
        form = MoviesForm()
    return render(request, 'add.html', {'form': form})


def edit_movie(request, movie_id):
    movie = get_object_or_404(Movies, movie_id=movie_id)
    if request.method == 'POST':
        form = MoviesForm(request.POST, instance=movie)
        if form.is_valid():
            form.save()
            return redirect('movies')
    else:
        form = MoviesForm(instance=movie)
    return render(request, 'edit.html', {'form': form, 'movie': movie})


def delete_movie(request, movie_id):
    movie = get_object_or_404(Movies, movie_id=movie_id)
    if request.method == 'POST':
        movie.delete()
        return redirect('movies')
    return render(request, 'delete.html', {'redirect_url': reverse('movies')})


def positions(request):
    sort_by = request.GET.get('sort_by', 'position_id')
    sort_order = request.GET.get('sort_order', 'asc')

    columns = [
        {'field': 'position_id', 'label': 'ID'},
        {'field': 'title', 'label': 'Название'},
    ]
    for col in columns:
        if col['field'] == sort_by:
            col['sort_order'] = 'desc' if sort_order == 'asc' else 'asc'
        else:
            col['sort_order'] = 'asc'

    position_list = Positions.objects.all().order_by(f"{'-' if sort_order == 'desc' else ''}{sort_by}")

    context = {
        'positions': position_list,
        'columns': columns,
        'sort_by': sort_by,
        'sort_order': sort_order,
    }
    return render(request, 'positions/positions.html', context)


def position_detail(request, position_id):
    position = get_object_or_404(Positions, position_id=position_id)
    return render(request, 'positions/position_detail.html', {'position': position})


def add_position(request):
    if request.method == 'POST':
        form = PositionsForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('positions')
    else:
        form = PositionsForm()
    return render(request, 'add.html', {'form': form})


def edit_position(request, position_id):
    position = get_object_or_404(Positions, position_id=position_id)
    if request.method == 'POST':
        form = PositionsForm(request.POST, instance=position)
        if form.is_valid():
            form.save()
            return redirect('positions')
    else:
        form = PositionsForm(instance=position)
    return render(request, 'edit.html', {'form': form, 'position': position})


def delete_position(request, position_id):
    position = get_object_or_404(Positions, position_id=position_id)
    if request.method == 'POST':
        position.delete()
        return redirect('positions')
    return render(request, 'delete.html', {'redirect_url': reverse('positions')})


def session_types(request):
    sort_by = request.GET.get('sort_by', 'session_type_id')
    sort_order = request.GET.get('sort_order', 'asc')

    columns = [
        {'field': 'session_type_id', 'label': 'ID'},
        {'field': 'name', 'label': 'Название'},
    ]
    for col in columns:
        if col['field'] == sort_by:
            col['sort_order'] = 'desc' if sort_order == 'asc' else 'asc'
        else:
            col['sort_order'] = 'asc'

    session_type_list = SessionTypes.objects.all().order_by(f"{'-' if sort_order == 'desc' else ''}{sort_by}")

    context = {
        'session_types': session_type_list,
        'columns': columns,
        'sort_by': sort_by,
        'sort_order': sort_order,
    }
    return render(request, 'session_types/session_types.html', context)


def session_type_detail(request, session_type_id):
    session_type = get_object_or_404(SessionTypes, session_type_id=session_type_id)
    return render(request, 'session_types/session_type_detail.html', {'session_type': session_type})


def add_session_type(request):
    if request.method == 'POST':
        form = SessionTypesForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('session_types')
    else:
        form = SessionTypesForm()
    return render(request, 'add.html', {'form': form})


def edit_session_type(request, session_type_id):
    session_type = get_object_or_404(SessionTypes, session_type_id=session_type_id)
    if request.method == 'POST':
        form = SessionTypesForm(request.POST, instance=session_type)
        if form.is_valid():
            form.save()
            return redirect('session_types')
    else:
        form = SessionTypesForm(instance=session_type)
    return render(request, 'edit.html', {'form': form, 'session_type': session_type})


def delete_session_type(request, session_type_id):
    session_type = get_object_or_404(SessionTypes, session_type_id=session_type_id)
    if request.method == 'POST':
        session_type.delete()
        return redirect('session_types')
    return render(request, 'delete.html', {'redirect_url': reverse('session_types')})


def staff(request):
    sort_by = request.GET.get('sort_by', 'staff_id')
    sort_order = request.GET.get('sort_order', 'asc')

    columns = [
        {'field': 'staff_id', 'label': 'ID'},
        {'field': 'first_name', 'label': 'ФИО'},
        {'field': 'position', 'label': 'Должность'},
        {'field': 'phone', 'label': 'Телефон'},
    ]
    for col in columns:
        if col['field'] == sort_by:
            col['sort_order'] = 'desc' if sort_order == 'asc' else 'asc'
        else:
            col['sort_order'] = 'asc'

    staff_list = Staff.objects.all().order_by(f"{'-' if sort_order == 'desc' else ''}{sort_by}")
    paginator = Paginator(staff_list, 25)
    page_number = request.GET.get('page', 1)
    try:
        page_obj = paginator.get_page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)

    context = {
        'staff': page_obj,
        'columns': columns,
        'sort_by': sort_by,
        'sort_order': sort_order,
        'page_number': int(page_number)
    }
    return render(request, 'staff/staff.html', context)


def staff_detail(request, staff_id):
    staff_member = get_object_or_404(Staff, staff_id=staff_id)
    return render(request, 'staff/staff_detail.html', {'staff': staff_member})


def add_staff(request):
    if request.method == 'POST':
        form = StaffForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('staff')
    else:
        form = StaffForm()
    return render(request, 'add.html', {'form': form})


def edit_staff(request, staff_id):
    staff_member = get_object_or_404(Staff, staff_id=staff_id)
    if request.method == 'POST':
        form = StaffForm(request.POST, instance=staff_member)
        if form.is_valid():
            form.save()
            return redirect('staff')
    else:
        form = StaffForm(instance=staff_member)
    return render(request, 'edit.html', {'form': form, 'staff_member': staff_member})


def delete_staff(request, staff_id):
    staff_member = get_object_or_404(Staff, staff_id=staff_id)
    if request.method == 'POST':
        staff_member.delete()
        return redirect('staff')
    return render(request, 'delete.html', {'redirect_url': reverse('staff')})


def sessions(request):
    queryset = Sessions.objects.all()
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    session_type = request.GET.get('session_type')
    sort_by = request.GET.get('sort_by', 'session_date')
    sort_order = request.GET.get('sort_order', 'asc')

    if start_date and end_date and start_date > end_date:
        start_date, end_date = end_date, start_date
    if start_date:
        queryset = queryset.filter(session_date__gte=start_date)
    if end_date:
        queryset = queryset.filter(session_date__lte=end_date)
    if session_type:
        queryset = queryset.filter(session_type_id=session_type)

    columns = [
        {'field': 'session_id', 'label': 'ID'},
        {'field': 'session_date', 'label': 'Дата'},
        {'field': 'session_time', 'label': 'Время'},
        {'field': 'session_type__name', 'label': 'Тип сессии'},
        {'field': 'movie__title', 'label': 'Фильм'},
        {'field': 'hall__name', 'label': 'Зал'},
    ]
    for col in columns:
        if col['field'] == sort_by:
            col['sort_order'] = 'desc' if sort_order == 'asc' else 'asc'
        else:
            col['sort_order'] = 'asc'

    queryset = queryset.annotate(
        full_session_type=Concat('session_type__name', V(' '), 'movie__title', output_field=CharField())
    ).order_by(f"{'-' if sort_order == 'desc' else ''}{sort_by}")

    paginator = Paginator(queryset, 25)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.get_page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)

    context = {
        'sessions': page_obj,
        'session_types': SessionTypes.objects.all(),
        'start_date': start_date if start_date is not None else '',
        'end_date': end_date if start_date is not None else '',
        'selected_session_type': int(session_type) if session_type is not None and len(session_type) > 0 else '',
        'columns': columns,
        'sort_by': sort_by,
        'sort_order': sort_order,
        'page_number': page_number,
    }
    return render(request, 'sessions/sessions.html', context)


def session_detail(request, session_id):
    session = get_object_or_404(Sessions, session_id=session_id)
    return render(request, 'sessions/session_detail.html', {'session': session})


def add_session(request):
    if request.method == 'POST':
        form = SessionsForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('sessions')
    else:
        form = SessionsForm()
    return render(request, 'add.html', {'form': form})


def edit_session(request, session_id):
    session = get_object_or_404(Sessions, session_id=session_id)
    if request.method == 'POST':
        form = SessionsForm(request.POST, instance=session)
        if form.is_valid():
            form.save()
            return redirect('sessions')
    else:
        form = SessionsForm(instance=session)
    return render(request, 'edit.html', {'form': form, 'session': session})


def delete_session(request, session_id):
    session = get_object_or_404(Sessions, session_id=session_id)
    if request.method == 'POST':
        session.delete()
        return redirect('sessions')
    return render(request, 'delete.html', {'redirect_url': reverse('sessions')})


def tickets(request):
    queryset = Tickets.objects.all()
    sort_by = request.GET.get('sort_by', 'ticket_id')
    sort_order = request.GET.get('sort_order', 'asc')
    min_price = request.GET.get('min_price')
    max_price = request.GET.get('max_price')

    if min_price and max_price and min_price > max_price:
        min_price, max_price = max_price, min_price
    if min_price:
        queryset = queryset.filter(price__gte=min_price)
    if max_price:
        queryset = queryset.filter(price__lte=max_price)

    columns = [
        {'field': 'ticket_id', 'label': 'ID'},
        {'field': 'session', 'label': 'Сеанс'},
        {'field': 'price', 'label': 'Цена'},
        {'field': 'row_number', 'label': 'Ряд'},
        {'field': 'seat_number', 'label': 'Место'},
    ]
    for col in columns:
        if col['field'] == sort_by:
            col['sort_order'] = 'desc' if sort_order == 'asc' else 'asc'
        else:
            col['sort_order'] = 'asc'

    queryset = queryset.order_by(f"{'-' if sort_order == 'desc' else ''}{sort_by}")

    paginator = Paginator(queryset, 25)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.get_page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)

    context = {
        'tickets': page_obj,
        'min_price': min_price if min_price is not None else '',
        'max_price': max_price if max_price is not None else '',
        'columns': columns,
        'sort_by': sort_by,
        'sort_order': sort_order,
        'page_number': page_number,
    }
    return render(request, 'tickets/tickets.html', context)


def ticket_detail(request, ticket_id):
    ticket = get_object_or_404(Tickets, ticket_id=ticket_id)
    return render(request, 'tickets/ticket_detail.html', {'ticket': ticket})


def add_ticket(request):
    if request.method == 'POST':
        form = TicketsForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('tickets')
    else:
        form = TicketsForm()
    return render(request, 'add.html', {'form': form})


def edit_ticket(request, ticket_id):
    ticket = get_object_or_404(Tickets, ticket_id=ticket_id)
    if request.method == 'POST':
        form = TicketsForm(request.POST, instance=ticket)
        if form.is_valid():
            form.save()
            return redirect('tickets')
    else:
        form = TicketsForm(instance=ticket)
    return render(request, 'edit.html', {'form': form, 'ticket': ticket})


def delete_ticket(request, ticket_id):
    ticket = get_object_or_404(Tickets, ticket_id=ticket_id)
    if request.method == 'POST':
        ticket.delete()
        return redirect('tickets')
    return render(request, 'delete.html', {'ticket': ticket})


def sales(request):
    queryset = Sales.objects.all()
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    staff = request.GET.get('staff')
    customer = request.GET.get('customer')
    sort_by = request.GET.get('sort_by', 'ticket_id')
    sort_order = request.GET.get('sort_order', 'asc')

    if start_date and end_date and start_date > end_date:
        start_date, end_date = end_date, start_date
    if start_date:
        queryset = queryset.filter(date__gte=start_date)
    if end_date:
        queryset = queryset.filter(date__lte=end_date)
    if staff:
        queryset = queryset.filter(staff_id=staff)
    if customer:
        queryset = queryset.filter(customer_id=customer)

    columns = [
        {'field': 'sale_id', 'label': 'ID'},
        {'field': 'ticket', 'label': 'Билет'},
        {'field': 'staff', 'label': 'Сотрудник'},
        {'field': 'date', 'label': 'Дата'},
        {'field': 'payment_type', 'label': 'Тип оплаты'},
        {'field': 'customer', 'label': 'Покупатель'},
    ]
    for col in columns:
        if col['field'] == sort_by:
            col['sort_order'] = 'desc' if sort_order == 'asc' else 'asc'
        else:
            col['sort_order'] = 'asc'

    queryset = queryset.order_by(f"{'-' if sort_order == 'desc' else ''}{sort_by}")

    paginator = Paginator(queryset, 25)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.get_page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)

    context = {
        'sales': page_obj,
        'staffs': Staff.objects.all(),
        'customers': Customers.objects.all(),
        'start_date': start_date if start_date is not None else '',
        'end_date': end_date if end_date is not None else '',
        'selected_staff': int(staff) if staff is not None and len(staff) > 0 else '',
        'selected_customer': int(customer) if customer is not None and len(customer) > 0 else '',
        'columns': columns,
        'sort_by': sort_by,
        'sort_order': sort_order,
        'page_number': page_number,
    }
    return render(request, 'sales/sales.html', context)


def sale_detail(request, sale_id):
    sale = get_object_or_404(Sales, sale_id=sale_id)
    return render(request, 'sales/sale_detail.html', {'sale': sale})


def add_sale(request):
    if request.method == 'POST':
        form = SalesForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('sales')
    else:
        form = SalesForm()
    return render(request, 'add.html', {'form': form})


def edit_sale(request, sale_id):
    sale = get_object_or_404(Sales, sale_id=sale_id)
    if request.method == 'POST':
        form = SalesForm(request.POST, instance=sale)
        if form.is_valid():
            form.save()
            return redirect('sales')
    else:
        form = SalesForm(instance=sale)
    return render(request, 'edit.html', {'form': form, 'sale': sale})


def delete_sale(request, sale_id):
    sale = get_object_or_404(Sales, sale_id=sale_id)
    if request.method == 'POST':
        sale.delete()
        return redirect('sales')
    return render(request, 'delete.html', {'sale': sale})


def sales_report(request):
    queryset = Sales.objects.all()
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    staff = request.GET.get('staff')
    customer = request.GET.get('customer')

    if start_date and end_date and start_date > end_date:
        start_date, end_date = end_date, start_date
    if start_date:
        queryset = queryset.filter(date__gte=start_date)
    if end_date:
        queryset = queryset.filter(date__lte=end_date)
    if staff:
        queryset = queryset.filter(staff_id=staff)
    if customer:
        queryset = queryset.filter(customer_id=customer)

    if request.method == 'POST' and 'export_excel' in request.POST:
        def background_task():
            export_to_excel(queryset, 'sales')
        thread = threading.Thread(target=background_task)
        thread.start()
        messages.success(request, 'Отчет экспортирован')
        return redirect('report')

    paginator = Paginator(queryset, 25)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.get_page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)

    context = {
        'sales': page_obj,
        'staffs': Staff.objects.all(),
        'customers': Customers.objects.all(),
        'start_date': start_date,
        'end_date': end_date,
        'selected_staff': int(staff) if staff is not None and len(staff) > 0 else None,
        'selected_customer': int(customer) if customer is not None and len(customer) > 0 else None,
    }
    return render(request, 'report/report.html', context)


def staff_report(request):
    queryset = Staff.objects.annotate(
        total_sales=Count('sales')
    ).order_by('-total_sales')
    staff = request.GET.get('staff')
    if staff:
        queryset = queryset.filter(staff_id=staff)

    if request.method == 'POST' and 'export_excel' in request.POST:
        def background_task():
            export_to_excel(queryset, 'staff')
        thread = threading.Thread(target=background_task)
        thread.start()
        messages.success(request, 'Отчет экспортирован')
        return redirect('staff-report')

    paginator = Paginator(queryset, 25)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.get_page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)

    context = {
        'employees': page_obj,
        'all_staff': Staff.objects.all(),
        'selected_staff': int(staff) if staff is not None and len(staff) > 0 else None
    }
    return render(request, 'report/staff_report.html', context)


def movies_report(request):
    queryset = Movies.objects.annotate(
        total_tickets_sold=Count('sessions__tickets')
    ).order_by('-total_tickets_sold')

    movie_id = request.GET.get('movie')
    genre = request.GET.get('genre')

    if movie_id:
        queryset = queryset.filter(movie_id=movie_id)
    if genre:
        queryset = queryset.filter(genre=genre)

    if request.method == 'POST' and 'export_excel' in request.POST:
        def background_task():
            export_to_excel(queryset, 'movies')
        thread = threading.Thread(target=background_task)
        thread.start()
        messages.success(request, 'Отчет экспортирован')

        return redirect('movies-report')

    paginator = Paginator(queryset, 25)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.get_page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)

    unique_genres = Movies.objects.order_by(Lower('genre')).values_list('genre', flat=True).distinct()
    context = {
        'movies': page_obj,
        'all_movies': Movies.objects.all(),
        'all_genres': unique_genres,
        'selected_movie': int(movie_id) if movie_id is not None and len(movie_id) > 0 else None,
        'selected_genre': genre if genre is not None and len(genre) > 0 else None
    }
    return render(request, 'report/movies_report.html', context)


def export_to_excel(queryset, report_type):
    workbook = openpyxl.Workbook()
    sheet = workbook.active

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    if report_type == 'sales':
        columns = [
            'ID', 'Дата', 'Сотрудник', 'Покупатель', 'Стоимость', 'Дата и время сеанса'
        ]
        sheet.append(columns)

        doc.add_heading('Отчет по продажам', 0)
        table = doc.add_table(rows=1, cols=len(columns))
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(columns):
            hdr_cells[i].text = column_name

        for sale in queryset:
            staff_name = f"{sale.staff.first_name} {sale.staff.last_name}"
            customer_name = f"{sale.customer.first_name} {sale.customer.last_name}"
            session_datetime = f"{sale.ticket.session.session_date} {sale.ticket.session.session_time}"

            sheet.append([
                sale.sale_id, sale.date, staff_name, customer_name, sale.ticket.price, session_datetime
            ])
            row_cells = table.add_row().cells
            for i, cell_data in enumerate([sale.sale_id, sale.date, staff_name, customer_name, sale.ticket.price, session_datetime]):
                row_cells[i].text = str(cell_data)

        filename = f"sales_report_{now().strftime('%Y%m%d_%H%M%S')}"

    elif report_type == 'staff':
        columns = ['ID', 'Имя', 'Фамилия', 'Отчество', 'Должность', 'Кол-во продаж']
        sheet.append(columns)

        doc.add_heading('Отчет по сотрудникам', 0)

        table = doc.add_table(rows=1, cols=len(columns))
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(columns):
            hdr_cells[i].text = column_name

        for staff in queryset:
            if staff.position is None:
                data = [
                    staff.staff_id, staff.first_name, staff.last_name, staff.middle_name, None, staff.total_sales
                ]

            else:
                data = [
                    staff.staff_id, staff.first_name, staff.last_name, staff.middle_name, staff.position.title, staff.total_sales
                ]

            sheet.append(data)
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(data):
                row_cells[i].text = str(cell_data)

        filename = f"staff_report_{now().strftime('%Y%m%d_%H%M%S')}"

    elif report_type == 'movies':
        columns = ['ID', 'Название', 'Жанр', 'Длительность', 'Рейтинг', 'Кол-во билетов']
        sheet.append(columns)

        doc.add_heading('Отчет по фильмам', 0)
        table = doc.add_table(rows=1, cols=len(columns))
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(columns):
            hdr_cells[i].text = column_name

        for movie in queryset:
            data = [
                movie.movie_id, movie.title, movie.genre, movie.duration, movie.rating,
                movie.total_tickets_sold
            ]
            sheet.append(data)
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(
                    [movie.movie_id, movie.title, movie.genre, movie.duration, movie.rating, movie.total_tickets_sold]):
                row_cells[i].text = str(cell_data)

        filename = f"movies_report_{now().strftime('%Y%m%d_%H%M%S')}"

    filepath = os.path.join(settings.MEDIA_ROOT, 'reports', filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    workbook.save(f'{filepath}.xlsx')
    doc.save(f'{filepath}.docx')

    return redirect('report')


def get_rows(request):
    session_id = request.GET.get('session_id')
    session = Sessions.objects.get(pk=session_id)
    hall_capacity = session.hall.capacity
    rows = list(range(1, hall_capacity // 20 + 1))

    return JsonResponse({'rows': rows})
